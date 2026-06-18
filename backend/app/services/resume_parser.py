from pathlib import Path

import fitz
from docx import Document


class ResumeParseError(Exception):
    pass


def extract_resume_text(file_path: str, file_type: str | None) -> str:
    path = Path(file_path)
    if not path.exists():
        raise ResumeParseError("简历文件不存在")

    suffix = (file_type or path.suffix.lstrip(".")).lower()
    if suffix == "pdf":
        text = _extract_pdf_text(path)
    elif suffix == "docx":
        text = _extract_docx_text(path)
    elif suffix == "doc":
        raise ResumeParseError("暂不支持解析 .doc 格式，请上传 .pdf 或 .docx 简历")
    else:
        raise ResumeParseError("不支持的简历文件格式")

    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not text:
        raise ResumeParseError("未能从简历中提取到有效文本")
    return text


def _extract_pdf_text(path: Path) -> str:
    try:
        with fitz.open(path) as doc:
            return "\n".join(page.get_text("text") for page in doc)
    except Exception as exc:
        raise ResumeParseError("PDF 简历解析失败") from exc


def _extract_docx_text(path: Path) -> str:
    try:
        document = Document(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        tables = []
        for table in document.tables:
            for row in table.rows:
                tables.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs + tables)
    except Exception as exc:
        raise ResumeParseError("DOCX 简历解析失败") from exc
